import fitz  # PyMuPDF
import re
from datetime import datetime
from pathlib import Path
from typing import Optional
from dataclasses import dataclass
from loguru import logger
import threading

from .layout_parser import reorder_page_text, detect_layout_stats
from .image_ocr import (
    get_ocr_engine,
    ocr_image_bytes,
    extract_pdf_page_image_text,
    extract_docx_image_text,
    extract_docx_table_text,
)
from .metadata_quality import (
    clean_authors,
    display_title,
    humanize_filename,
    is_poor_title,
    normalize_ocr_page_text,
    resolve_paper_title,
    strip_uuid_prefix,
)

_ocr_lock = threading.Lock()

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm", ".epub",
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


@dataclass
class ExtractedDocument:
    path: str
    filename: str
    title: str
    authors: str  # JSON array string
    year: Optional[int]
    doi: str
    page_count: int
    file_size: int
    language: str
    text_by_page: dict[int, str]  # page_number -> text
    full_text: str
    ocr_pages_count: int = 0
    ocr_pages_failed: int = 0
    is_scanned: bool = False
    layout_stats: Optional[dict] = None
    suggested_title: Optional[str] = None


def create_image_stub_document(file_path: str, original_filename: str | None = None) -> ExtractedDocument:
    """Fast metadata-only stub for image imports; OCR runs in background."""
    path = Path(file_path)
    name = original_filename or path.name
    title = resolve_paper_title(filename=name, stored_path=str(path))
    stub_text = f"[Image: {name}]\n(OCR and indexing in progress...)"
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=name,
        title=title,
        authors="[]",
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language="vi",
        text_by_page={1: stub_text},
        full_text=stub_text,
        ocr_pages_count=0,
        ocr_pages_failed=0,
        is_scanned=True,
        suggested_title=title,
    )


def _extract_title_from_text(text: str, fallback: str, max_chars: int = 150) -> str:
    """Get first meaningful line from text as a suggested title."""
    for line in text.split("\n"):
        line = normalize_ocr_page_text(line).strip()
        if not line:
            continue
        # Skip short lines (page numbers, headers)
        if len(line) < 10:
            continue
        # Skip lines that are mostly numbers or symbols
        alpha_ratio = sum(1 for c in line if c.isalpha()) / max(len(line), 1)
        if alpha_ratio < 0.4:
            continue
        # Skip logo/caption/junk lines that pollute library titles
        if is_poor_title(line):
            continue
        return line[:max_chars].strip()
    cleaned_fallback = humanize_filename(fallback) or fallback
    return cleaned_fallback[:max_chars].strip()


def extract_document(file_path: str) -> Optional[ExtractedDocument]:
    path = Path(file_path)
    if not path.exists():
        return None

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".md":
        return _extract_markdown(file_path)
    elif ext in (".html", ".htm"):
        return _extract_html(file_path)
    elif ext == ".epub":
        return _extract_epub(file_path)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_image(file_path)
    return None


def extract_pdf(file_path: str) -> Optional[ExtractedDocument]:
    return _extract_pdf(file_path)


def _process_single_page(file_path: str, page_num: int, collect_layout: bool = False, include_image_ocr: bool = False) -> tuple[int, str, bool, bool, Optional[dict]]:
    import fitz
    import re
    
    text = ""
    ocr_attempted = False
    ocr_succeeded = False
    layout = None
    try:
        doc = fitz.open(file_path)
        page = doc[page_num]
        text = page.get_text("text")

        # Multi-column layout reordering (safe fallback to get_text if single-column)
        if len(text.strip()) > 100:
            reordered = reorder_page_text(page)
            if reordered != text:
                if collect_layout:
                    layout = detect_layout_stats(page)
                    logger.info(f"Page {page_num+1}: detected {layout.get('columns', '?')}-column layout, reordered text")
                text = reordered
        
        # Detect text with corrupted character encoding (control chars in wrong ranges)
        bad_chars = sum(1 for c in text if ord(c) < 0x09 or 0x0E <= ord(c) < 0x20 or 0x80 <= ord(c) < 0xA0)
        is_garbled = bad_chars > max(3, len(text.strip()) * 0.05)
        if is_garbled and len(text.strip()) > 0:
            text_html = page.get_text("html")
            if text_html:
                html_text = re.sub(r'<[^>]+>', ' ', text_html)
                html_bad = sum(1 for c in html_text if ord(c) < 0x09 or 0x0E <= ord(c) < 0x20 or 0x80 <= ord(c) < 0xA0)
                if html_bad < bad_chars:
                    text = html_text
                    is_garbled = False

        # Detect text that looks valid but is semantically garbage (wrong font mapping)
        if not is_garbled and len(text.strip()) >= 40:
            words = [w for w in text.split() if w.strip()]
            if words:
                alpha_ratio = sum(1 for c in text if c.isalpha()) / max(len(text.strip()), 1)
                non_alpha_words = sum(1 for w in words if sum(1 for c in w if c.isalpha()) / max(len(w), 1) < 0.3)
                garbage_word_ratio = non_alpha_words / max(len(words), 1)
                # Heuristic: low alphabetic ratio + many "words" that are mostly punctuation = garbled
                if alpha_ratio < 0.4 and garbage_word_ratio > 0.5:
                    is_garbled = True

        if is_garbled or len(text.strip()) < 40:
            ocr_attempted = True
            try:
                ocr_engine = get_ocr_engine()
                logger.info(f"Page {page_num + 1} appears to be a scanned page (text length: {len(text.strip())}). Running OCR...")
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")

                with _ocr_lock:
                    ocr_results, elapse = ocr_engine(img_bytes)

                if ocr_results:
                    ocr_text_list = [res[1] for res in ocr_results if res and len(res) > 1]
                    ocr_text = "\n".join(ocr_text_list)
                    if ocr_text.strip():
                        text = normalize_ocr_page_text(ocr_text)
                        ocr_succeeded = True
                        total_elapse = sum(elapse) if isinstance(elapse, (list, tuple)) else (elapse or 0.0)
                        logger.info(f"Page {page_num + 1} OCR completed in {total_elapse:.2f}s")
            except Exception as e:
                logger.warning(f"OCR failed for page {page_num + 1}: {e}")

        # Embedded figure OCR is opt-in. Running OCR over every image during
        # import can stall PDF parsing on image-heavy papers.
        if include_image_ocr:
            image_snippets = extract_pdf_page_image_text(
                page,
                doc,
                page_num=page_num + 1,
                page_text=text,
            )
            if image_snippets:
                text = (text.rstrip() + "\n\n" + "\n".join(image_snippets)).strip()

        if text:
            text = normalize_ocr_page_text(text)
        doc.close()
    except Exception as e:
        logger.error(f"Error processing page {page_num + 1} from {file_path}: {e}")
        
    ocr_failed = ocr_attempted and not ocr_succeeded
    return page_num, text, ocr_attempted, ocr_failed, layout


def _clean_metadata_string(text: str) -> str:
    if not text:
        return ""
    try:
        # Try encoding latin-1 to byte representation and decode to UTF-8
        decoded = text.encode("latin-1").decode("utf-8")
        # Check if it contains common Vietnamese accented characters
        if any(c in decoded for c in "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ"):
            return decoded
        return text
    except (UnicodeEncodeError, UnicodeDecodeError):
        return text


def _serialize_authors(authors_list: list) -> str:
    import json
    return json.dumps(authors_list, ensure_ascii=False)


def _parse_pdf_authors(raw: str) -> list[str]:
    cleaned = _clean_metadata_string(raw).strip()
    if not cleaned:
        return []
    parts = re.split(r"\s*(?:;|\r?\n|\band\b|\|)\s*", cleaned, flags=re.IGNORECASE)
    if len(parts) == 1 and "," in cleaned:
        parts = [part.strip() for part in cleaned.split(",")]
    return clean_authors([part.strip() for part in parts if part.strip()])


def _extract_metadata_year(*values: str) -> Optional[int]:
    for value in values:
        if not value:
            continue
        match = re.search(r"(?:D:)?((?:19|20)\d{2})", str(value))
        if match:
            year = int(match.group(1))
            if 1900 <= year <= datetime.now().year + 1:
                return year
    return None


def _extract_pdf(file_path: str) -> Optional[ExtractedDocument]:
    path = Path(file_path)

    try:
        doc = fitz.open(file_path)
    except Exception:
        return None

    page_count = len(doc)
    doc.close()

    text_by_page: dict[int, str] = {}
    ocr_pages_count = 0
    ocr_pages_failed = 0
    layout_stats: dict = {}

    # Parse pages sequentially. PyMuPDF plus OCR inside a thread pool can leave
    # one worker stuck and keep the whole import at "parsing 25%" indefinitely.
    try:
        for page_num in range(page_count):
            p_num, text, ocr_attempted, ocr_failed, layout = _process_single_page(
                file_path,
                page_num,
                collect_layout=True,
                include_image_ocr=False,
            )
            text_by_page[p_num + 1] = text
            if ocr_attempted:
                ocr_pages_count += 1
            if ocr_failed:
                ocr_pages_failed += 1
            if layout:
                layout_stats[p_num + 1] = layout
    except Exception as e:
        logger.error(f"Error during PDF parsing: {e}")
        return None

    full_text_parts = [text_by_page[p] for p in sorted(text_by_page.keys())]
    full_text = "\n".join(full_text_parts)

    # Prefer humanized original name (strip storage UUID) over raw path.stem
    filename_fallback = humanize_filename(path.name) or strip_uuid_prefix(path.stem) or path.stem

    try:
        doc = fitz.open(file_path)
        meta = doc.metadata or {}
        metadata_title = _clean_metadata_string(meta.get("title", "").strip())
        authors_raw = _clean_metadata_string(meta.get("author", "").strip())
        authors_list = _parse_pdf_authors(authors_raw)
        authors_json = _serialize_authors(authors_list)

        year = _extract_metadata_year(
            meta.get("creationDate", ""),
            meta.get("modDate", ""),
            filename_fallback,
            path.stem,
        )

        doi = meta.get("identifier", "").strip() or ""
        doc.close()
    except Exception as meta_err:
        logger.warning(f"Failed to extract metadata for {file_path}: {meta_err}")
        metadata_title = ""
        authors_json = "[]"
        year = None
        doi = ""

    language = _detect_language(full_text[:2000])

    # Suggested title from first page — only used if it passes quality checks
    first_page_text = text_by_page.get(1, "")
    suggested_title = _extract_title_from_text(first_page_text, filename_fallback)

    title = resolve_paper_title(
        metadata_title=metadata_title,
        suggested_title=suggested_title,
        filename=path.name,
        stored_path=str(path),
    )

    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors=authors_json,
        year=year,
        doi=doi,
        page_count=len(text_by_page),
        file_size=path.stat().st_size,
        language=language,
        text_by_page=text_by_page,
        full_text=full_text,
        ocr_pages_count=ocr_pages_count,
        ocr_pages_failed=ocr_pages_failed,
        is_scanned=ocr_pages_count > 0,
        layout_stats=layout_stats if layout_stats else None,
        suggested_title=suggested_title,
    )


def _make_suggested_title(doc: ExtractedDocument) -> str:
    """Compute suggested_title from first page text if not already set."""
    if doc.suggested_title:
        return doc.suggested_title
    first_text = doc.text_by_page.get(1, "")
    return _extract_title_from_text(first_text, doc.filename)


def _extract_docx(file_path: str) -> Optional[ExtractedDocument]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed")
        return None

    path = Path(file_path)
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        return None

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    table_parts = extract_docx_table_text(doc)
    image_parts = extract_docx_image_text(doc)

    full_text_parts = paragraphs + table_parts + image_parts
    full_text = "\n\n".join(full_text_parts)
    if not full_text.strip():
        return None

    # Try to get title from first heading or first line
    title = path.stem
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            title = para.text.strip()
            break
    if title == path.stem and paragraphs:
        title = paragraphs[0]

    authors_list = []
    author_para = None
    for para in doc.paragraphs:
        if para.text.strip().lower().startswith("author"):
            author_para = para.text.strip()
            break
    if author_para:
        author_str = author_para.split(":", 1)[-1].strip()
        authors_list = [a.strip() for a in author_str.split(";") if a.strip()]

    language = _detect_language(full_text[:2000])

    suggested = _extract_title_from_text(full_text, path.stem)
    has_image_ocr = bool(image_parts)
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors=_serialize_authors(authors_list),
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: full_text},
        full_text=full_text,
        ocr_pages_count=1 if has_image_ocr else 0,
        ocr_pages_failed=0,
        is_scanned=has_image_ocr,
        suggested_title=suggested,
    )


def _extract_txt(file_path: str) -> Optional[ExtractedDocument]:
    path = Path(file_path)
    try:
        raw = path.read_bytes()
        full_text = raw.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Failed to read TXT {file_path}: {e}")
        return None

    if not full_text.strip():
        return None

    lines = [l.strip() for l in full_text.splitlines() if l.strip()]
    title = lines[0] if lines else path.stem

    language = _detect_language(full_text[:2000])

    suggested = _extract_title_from_text(full_text, path.stem)
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors="[]",
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: full_text},
        full_text=full_text,
        ocr_pages_count=0,
        ocr_pages_failed=0,
        is_scanned=False,
        suggested_title=suggested,
    )


def _extract_markdown(file_path: str) -> Optional[ExtractedDocument]:
    path = Path(file_path)
    try:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Failed to read MD {file_path}: {e}")
        return None

    if not text.strip():
        return None

    # Try to extract YAML frontmatter
    title = path.stem
    authors_list = []
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            front = text[3:end].strip()
            text = text[end + 3:].strip()
            for line in front.splitlines():
                if ":" in line:
                    key, val = line.split(":", 1)
                    key = key.strip().lower()
                    val = val.strip().strip('"').strip("'")
                    if key == "title" and val:
                        title = val
                    elif key == "author" or key == "authors":
                        authors_list = [a.strip() for a in val.split(",") if a.strip()]

    # Remove markdown syntax for clean plain text; keep image alt text
    clean = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[Image: \1]', text)
    clean = re.sub(r'\[([^\]]*)\]\([^)]+\)', r'\1', clean)
    clean = re.sub(r'#{1,6}\s+', '', clean)
    clean = re.sub(r'[*_~`]', '', clean)
    clean = re.sub(r'^\s*[-*+]\s+', '', clean, flags=re.MULTILINE)
    clean = re.sub(r'^\s*\d+\.\s+', '', clean, flags=re.MULTILINE)

    # OCR local images referenced in markdown (same folder as the .md file)
    md_images = re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', text)
    image_ocr_parts: list[str] = []
    for alt, src in md_images:
        src = src.strip().split()[0]  # drop optional title after space
        if src.startswith(("http://", "https://", "data:")):
            if alt.strip():
                image_ocr_parts.append(f"[Image: {alt.strip()}]")
            continue
        img_path = (path.parent / src).resolve()
        if not img_path.exists() or img_path.suffix.lower() not in IMAGE_EXTENSIONS:
            if alt.strip():
                image_ocr_parts.append(f"[Image: {alt.strip()}]")
            continue
        try:
            img_bytes = img_path.read_bytes()
            ocr_text = ocr_image_bytes(img_bytes, skip_byte_size_check=True)
            label = alt.strip() or img_path.name
            if ocr_text:
                image_ocr_parts.append(f"[Image content {label}]: {ocr_text}")
            elif label:
                image_ocr_parts.append(f"[Image: {label}]")
        except Exception as exc:
            logger.debug(f"Skip markdown image {img_path}: {exc}")

    if image_ocr_parts:
        clean = clean.rstrip() + "\n\n" + "\n".join(image_ocr_parts)

    language = _detect_language(clean[:2000])

    suggested = _extract_title_from_text(clean, path.stem)
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors=_serialize_authors(authors_list),
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: clean},
        full_text=clean,
        ocr_pages_count=0,
        ocr_pages_failed=0,
        is_scanned=False,
        suggested_title=suggested,
    )


def _extract_html(file_path: str) -> Optional[ExtractedDocument]:
    path = Path(file_path)
    try:
        from lxml import html
        raw = path.read_bytes()
        doc = html.fromstring(raw)
    except Exception as e:
        logger.error(f"Failed to parse HTML {file_path}: {e}")
        return None

    # Remove script/style
    for el in doc.xpath("//script | //style | //nav | //footer | //header"):
        el.getparent().remove(el)

    title_el = doc.find(".//title")
    title = title_el.text_content().strip() if title_el is not None else path.stem

    body = doc.find(".//body")
    full_text = body.text_content().strip() if body is not None else doc.text_content().strip()

    image_notes: list[str] = []
    for img in doc.xpath("//img"):
        alt = (img.get("alt") or "").strip()
        title = (img.get("title") or "").strip()
        src = (img.get("src") or "").strip()
        if alt or title:
            image_notes.append(f"[Image: {alt or title}]")
        elif src:
            image_notes.append(f"[Image: {Path(src).name}]")

        if src and not src.startswith(("http://", "https://", "data:")):
            img_path = (path.parent / src).resolve()
            if img_path.exists() and img_path.suffix.lower() in IMAGE_EXTENSIONS:
                try:
                    ocr_text = ocr_image_bytes(img_path.read_bytes(), skip_byte_size_check=True)
                    if ocr_text:
                        label = alt or title or img_path.name
                        image_notes.append(f"[Image content {label}]: {ocr_text}")
                except Exception as exc:
                    logger.debug(f"Skip HTML image {img_path}: {exc}")

    if not full_text and not image_notes:
        return None

    if image_notes:
        full_text = (full_text.rstrip() + "\n\n" + "\n".join(image_notes)).strip() if full_text else "\n".join(image_notes)

    # Collapse whitespace
    full_text = re.sub(r'\s+', '\n', full_text).strip()

    language = _detect_language(full_text[:2000])

    suggested = _extract_title_from_text(full_text, path.stem)
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors="[]",
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: full_text},
        full_text=full_text,
        ocr_pages_count=0,
        ocr_pages_failed=0,
        is_scanned=False,
        suggested_title=suggested,
    )


def _extract_epub(file_path: str) -> Optional[ExtractedDocument]:
    try:
        from ebooklib import epub
    except ImportError:
        logger.error("ebooklib not installed")
        return None

    path = Path(file_path)
    try:
        book = epub.read_epub(file_path)
    except Exception as e:
        logger.error(f"Failed to parse EPUB {file_path}: {e}")
        return None

    text_parts = []
    title = path.stem
    authors_list = []

    # Get metadata
    title_meta = book.get_metadata("DC", "title")
    if title_meta:
        title = title_meta[0][0]

    authors_meta = book.get_metadata("DC", "creator")
    if authors_meta:
        authors_list = [a[0] for a in authors_meta]

    # Extract text from all items
    for item in book.get_items():
        if item.get_type() == 9:  # ITEM_DOCUMENT
            try:
                content = item.get_content()
                if content:
                    text = content.decode("utf-8", errors="replace")
                    # Strip HTML tags
                    text = re.sub(r'<[^>]+>', ' ', text)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if text:
                        text_parts.append(text)
            except Exception:
                continue

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        return None

    language = _detect_language(full_text[:2000])

    suggested = _extract_title_from_text(full_text, path.stem)
    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors=_serialize_authors(authors_list),
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: full_text},
        full_text=full_text,
        ocr_pages_count=0,
        ocr_pages_failed=0,
        is_scanned=False,
        suggested_title=suggested,
    )


def _extract_image(file_path: str) -> Optional[ExtractedDocument]:
    """Import a standalone image file via OCR."""
    path = Path(file_path)
    try:
        img_bytes = path.read_bytes()
    except Exception as exc:
        logger.error(f"Failed to read image {file_path}: {exc}")
        return None

    ocr_text = ocr_image_bytes(img_bytes, skip_byte_size_check=True)
    ocr_failed = 0
    if not ocr_text:
        logger.warning(f"No OCR text extracted from image: {file_path}")
        ocr_text = (
            f"[Image: {path.name}]\n"
            "(Could not extract text automatically. Use \"Re-run OCR\" in the library.)"
        )
        ocr_failed = 1
    else:
        ocr_text = normalize_ocr_page_text(ocr_text)

    suggested = _extract_title_from_text(ocr_text, humanize_filename(path.name) or path.name)
    title = resolve_paper_title(
        suggested_title=suggested,
        filename=path.name,
        stored_path=str(path),
    )
    language = _detect_language(ocr_text[:2000])

    return ExtractedDocument(
        path=str(path.absolute()),
        filename=path.name,
        title=title,
        authors="[]",
        year=None,
        doi="",
        page_count=1,
        file_size=path.stat().st_size,
        language=language,
        text_by_page={1: ocr_text},
        full_text=ocr_text,
        ocr_pages_count=0 if ocr_failed else 1,
        ocr_pages_failed=ocr_failed,
        is_scanned=True,
        suggested_title=suggested,
    )


def _detect_language(text: str) -> str:
    vietnamese_chars = set("ăâđêôơưàáãảạằắẳẵặầấẩẫậèéẻẽẹềếểễệìíỉĩịòóỏõọồốổỗộờớởỡợùúủũụừứửữựỳýỷỹỵ")
    count_vn = sum(1 for c in text.lower() if c in vietnamese_chars)
    if count_vn > 5:
        return "vi"
    return "en"
