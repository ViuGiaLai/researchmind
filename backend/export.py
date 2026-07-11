"""
ResearchMind VN — Paper Export Module

Endpoints:
- GET /api/papers/{paper_id}/export/html   → Full paper as styled HTML page
- GET /api/papers/{paper_id}/export/docx   → Full paper as Word document
"""

# Export display labels are intentionally in English — academic exports standardise on English as a neutral language.

import io
import json
import re
from datetime import datetime
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import HTMLResponse, StreamingResponse
from loguru import logger
from sqlalchemy.orm import Session

from common.i18n import t, get_language
from db.database import get_session
from db.models import Paper, Chunk

router = APIRouter(prefix="/api/papers", tags=["Export"])


# ─── Dependency ─────────────────────────────────────────────────

def _get_db(request: Request):
    """Get a DB session from the FastAPI app state engine, with auto-cleanup."""
    engine = request.app.state.engine
    session = get_session(engine)
    try:
        yield session
    finally:
        session.close()


# ─── Helpers ────────────────────────────────────────────────────

def _parse_authors(authors_str: str) -> list[str]:
    if not authors_str:
        return []
    try:
        val = json.loads(authors_str)
        if isinstance(val, list):
            return val
    except (json.JSONDecodeError, TypeError):
        pass
    
    try:
        val = json.loads(authors_str.replace("'", '"'))
        if isinstance(val, list):
            return val
    except Exception:
        pass
        
    import re
    cleaned = re.sub(r"[\[\]'\"#]", "", authors_str)
    return [a.strip() for a in cleaned.split(",") if a.strip()]

def _get_paper_data(paper_id: str, session: Session) -> dict | None:
    """Fetch paper metadata + chunks from DB."""
    paper = session.query(Paper).filter(Paper.id == paper_id).first()
    if not paper:
        return None

    authors_list = _parse_authors(paper.authors)

    # Parse tags
    try:
        tags_list = json.loads(paper.tags) if paper.tags else []
    except (json.JSONDecodeError, TypeError):
        tags_list = []

    # Get chunks ordered by index
    chunks = (
        session.query(Chunk)
        .filter(Chunk.paper_id == paper_id)
        .order_by(Chunk.chunk_index.asc())
        .all()
    )

    # Group chunks by page for cleaner display
    chunks_by_page: dict[int, list[str]] = {}
    for c in chunks:
        page = c.page_number or 1
        if page not in chunks_by_page:
            chunks_by_page[page] = []
        chunks_by_page[page].append(c.content)

    return {
        "id": paper.id,
        "title": paper.title or paper.filename.replace(".pdf", "").replace("_", " "),
        "filename": paper.filename,
        "authors": authors_list,
        "year": paper.year,
        "doi": paper.doi or "",
        "language": paper.language,
        "page_count": paper.page_count or 0,
        "file_size": paper.file_size,
        "tags": tags_list,
        "notes": paper.notes or "",
        "auto_summary": paper.auto_summary or "",
        "read_status": paper.read_status or "unread",
        "starred": bool(paper.starred),
        "status": paper.status,
        "chunks_by_page": chunks_by_page,
        "created_at": str(paper.created_at) if paper.created_at else "",
    }


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


# ═══════════════════════════════════════════════════════════════
# SHARED MARKDOWN PARSER — yields (event_type, data) tuples
# ═══════════════════════════════════════════════════════════════

def _parse_md_events(content: str):
    """
    Parse markdown content and yield (event_type, data) tuples.
    Each format handler (HTML/PDF/DOCX) consumes these events.
    """
    lines = content.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        stripped = lines[i].strip()

        # ── Code block ──
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            yield ("code_block", (code_lines, lang))
            continue

        # ── Table ──
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().split("|")]
                cells = [c for c in cells if c]
                table_rows.append(cells)
                i += 1
            yield ("table", (table_rows,))
            continue

        # ── Empty line ──
        if not stripped:
            i += 1
            continue

        # ── Heading 1 ──
        if stripped.startswith("# "):
            yield ("heading1", (stripped[2:],))
            i += 1
            continue

        # ── Heading 2 ──
        if stripped.startswith("## "):
            yield ("heading2", (stripped[3:],))
            i += 1
            continue

        # ── Heading 3 ──
        if stripped.startswith("### "):
            yield ("heading3", (stripped[4:],))
            i += 1
            continue

        # ── Horizontal rule ──
        if stripped.startswith("---"):
            yield ("hr", ())
            i += 1
            continue

        # ── Bullet list ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            items = []
            while i < n and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                items.append(lines[i].strip()[2:])
                i += 1
            yield ("bullet_list", (items,))
            continue

        # ── Numbered list ──
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                match = re.match(r"^\d+\.\s+(.+)$", lines[i].strip())
                if match:
                    items.append(match.group(1))
                i += 1
            yield ("numbered_list", (items,))
            continue

        # ── Blockquote ──
        if stripped.startswith("> ") or stripped == ">":
            quote_lines = []
            while i < n:
                s = lines[i].strip()
                if s.startswith("> "):
                    quote_lines.append(s[2:])
                elif s == ">":
                    quote_lines.append("")
                else:
                    break
                i += 1
            yield ("blockquote", (quote_lines,))
            continue

        # ── Paragraph (collects consecutive non-special lines) ──
        para_lines = []
        while i < n and lines[i].strip() and not lines[i].strip().startswith("#") \
                and not lines[i].strip().startswith("```") \
                and not lines[i].strip().startswith("|") \
                and not lines[i].strip().startswith("---") \
                and not lines[i].strip().startswith("> ") \
                and not lines[i].strip().startswith("- ") \
                and not lines[i].strip().startswith("* ") \
                and not re.match(r"^\d+\.\s+", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            yield ("paragraph", (para_lines,))


def _inline_html(text: str) -> str:
    """Convert inline markdown (bold, italic, code) to HTML."""
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", escaped)
    return escaped


def _md_to_html(md_text: str) -> str:
    """Convert Markdown to HTML with code blocks, tables, lists, headings, etc."""
    html_parts = []
    for event, data in _parse_md_events(md_text):
        if event == "code_block":
            code_lines, lang = data
            content = "\n".join(code_lines)
            escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
            lang_attr = f' class="language-{lang}"' if lang else ''
            html_parts.append(f"<pre><code{lang_attr}>{escaped}</code></pre>")
        elif event == "table":
            rows = data[0]
            if len(rows) < 2:
                continue
            header = rows[0]
            data_rows = rows[2:] if len(rows) > 2 else []
            thead = "<thead><tr>" + "".join(f"<th>{_inline_html(h)}</th>" for h in header) + "</tr></thead>"
            tbody = "<tbody>"
            for row in data_rows:
                tbody += "<tr>" + "".join(f"<td>{_inline_html(c)}</td>" for c in row) + "</tr>"
            tbody += "</tbody>"
            html_parts.append(f"<table>{thead}{tbody}</table>")
        elif event == "heading1":
            html_parts.append(f"<h1>{_inline_html(data[0])}</h1>")
        elif event == "heading2":
            html_parts.append(f"<h2>{_inline_html(data[0])}</h2>")
        elif event == "heading3":
            html_parts.append(f"<h3>{_inline_html(data[0])}</h3>")
        elif event == "hr":
            html_parts.append("<hr>")
        elif event == "bullet_list":
            items = "".join(f"<li>{_inline_html(item)}</li>" for item in data[0])
            html_parts.append(f"<ul>{items}</ul>")
        elif event == "numbered_list":
            items = "".join(f"<li>{_inline_html(item)}</li>" for item in data[0])
            html_parts.append(f"<ol>{items}</ol>")
        elif event == "blockquote":
            lines_data = data[0]
            content = "<br>".join(_inline_html(l) if l else "" for l in lines_data)
            html_parts.append(f"<blockquote>{content}</blockquote>")
        elif event == "paragraph":
            content = "<br>".join(_inline_html(l) for l in data[0])
            html_parts.append(f"<p>{content}</p>")
    return "\n".join(html_parts)


# ─── HTML Export ────────────────────────────────────────────────

@router.get("/{paper_id}/export/html")
async def export_paper_html(paper_id: str, db: Session = Depends(_get_db)):
    """Export full paper content as a styled HTML page."""
    data = _get_paper_data(paper_id, db)
    if not data:
        raise HTTPException(status_code=404, detail="Paper not found")

    title = _escape_html(data["title"])
    authors_str = ", ".join(data["authors"]) if data["authors"] else "Unknown"
    year_str = str(data["year"]) if data["year"] else "N/A"
    doi_str = _escape_html(data["doi"])
    tags_str = ", ".join(data["tags"]) if data["tags"] else "—"
    lang = data["language"].upper()
    pages = data["page_count"]

    # Build summary HTML
    summary_html = ""
    if data["auto_summary"]:
        summary_html = f"""
        <section class="section summary-section">
            <h2>🧠 AI Summary</h2>
            <div class="summary-content">{_md_to_html(data["auto_summary"])}</div>
        </section>"""

    # Build notes HTML
    notes_html = ""
    if data["notes"]:
        notes_html = f"""
        <section class="section notes-section">
            <h2>📝 Personal Notes</h2>
            <div class="notes-content">{_escape_html(data["notes"]).replace(chr(10), "<br>")}</div>
        </section>"""

    # Build content HTML by page
    content_html = ""
    for page_num in sorted(data["chunks_by_page"].keys()):
        chunks_text = data["chunks_by_page"][page_num]
        text_combined = "\n\n".join(chunks_text)
        content_html += f"""
        <section class="section page-section">
            <h2>📄 Page {page_num}</h2>
            <div class="page-content">{_escape_html(text_combined).replace(chr(10), "<br>")}</div>
        </section>"""

    if not content_html:
        content_html = '<p class="empty-state">No extracted text available for this paper.</p>'

    # Build citation (APA style)
    if data["authors"]:
        if len(data["authors"]) == 1:
            cite_author = data["authors"][0]
        elif len(data["authors"]) == 2:
            cite_author = f"{data['authors'][0]} & {data['authors'][1]}"
        else:
            cite_author = f"{data['authors'][0]} et al."
    else:
        cite_author = "Unknown"
    citation = f"{cite_author} ({year_str}). <em>{_escape_html(data['title'])}</em>."
    if doi_str:
        citation += f" <a href=\"https://doi.org/{doi_str}\">https://doi.org/{doi_str}</a>"

    today = datetime.now().strftime("%Y-%m-%d %H:%M")
    starred_badge = "⭐ Starred" if data["starred"] else ""

    html = f"""<!DOCTYPE html>
<html lang="{lang.lower()}">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} — ResearchMind VN</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'Georgia', 'Times New Roman', serif; font-size: 12pt; line-height: 1.8; color: #1a1a1a; background: #fafafa; padding: 0; }}
  .container {{ max-width: 900px; margin: 0 auto; background: #fff; padding: 48px 56px; min-height: 100vh; box-shadow: 0 0 40px rgba(0,0,0,0.06); }}
  @media (max-width: 768px) {{ .container {{ padding: 24px 20px; }} }}
  h1 {{ font-size: 26pt; font-weight: 700; color: #111; line-height: 1.3; margin-bottom: 8px; }}
  .meta-row {{ display: flex; flex-wrap: wrap; gap: 8px 24px; margin-bottom: 24px; padding-bottom: 20px; border-bottom: 2px solid #6366f1; font-size: 10.5pt; color: #555; }}
  .meta-row .meta-label {{ font-weight: 600; color: #333; }}
  .badge {{ display: inline-block; padding: 2px 10px; border-radius: 12px; font-size: 9pt; font-weight: 500; }}
  .badge-starred {{ background: #fef3c7; color: #92400e; }}
  .badge-lang {{ background: #e0e7ff; color: #3730a3; }}
  .badge-status {{ background: #d1fae5; color: #065f46; }}
  .section {{ margin-bottom: 32px; }}
  .section h2 {{ font-size: 16pt; font-weight: 700; color: #111; border-bottom: 1px solid #e5e7eb; padding-bottom: 8px; margin-bottom: 16px; }}
  .summary-content {{ background: #f5f3ff; border-left: 4px solid #6366f1; padding: 16px 20px; border-radius: 0 8px 8px 0; font-size: 11pt; line-height: 1.7; color: #333; }}
  .summary-content ul, .page-content ul {{ margin: 8px 0; padding-left: 24px; }}
  .summary-content li, .page-content li {{ margin-bottom: 4px; }}
  .summary-content h3 {{ font-size: 12pt; font-weight: 600; margin-top: 12px; margin-bottom: 6px; color: #333; }}
  .summary-content h4 {{ font-size: 11pt; font-weight: 600; margin-top: 10px; color: #444; }}
  .notes-content {{ background: #fffbeb; border-left: 4px solid #f59e0b; padding: 16px 20px; border-radius: 0 8px 8px 0; white-space: pre-wrap; font-size: 11pt; color: #333; }}
  .page-section {{ }}
  .page-content {{ font-size: 10.5pt; line-height: 1.8; color: #333; white-space: pre-wrap; background: #f9fafb; padding: 16px 20px; border-radius: 8px; border: 1px solid #e5e7eb; }}
  .page-content p {{ margin-bottom: 8px; }}
  .citation-section {{ background: #f0f9ff; border: 1px solid #bae6fd; border-radius: 8px; padding: 16px 20px; }}
  .citation-section h2 {{ border-bottom-color: #7dd3fc; }}
  .citation-text {{ font-size: 11pt; line-height: 1.7; color: #333; }}
  .citation-text a {{ color: #6366f1; text-decoration: none; }}
  .citation-text a:hover {{ text-decoration: underline; }}
  .empty-state {{ text-align: center; padding: 48px; color: #999; font-style: italic; }}
  .footer {{ margin-top: 48px; padding-top: 16px; border-top: 1px solid #e5e7eb; font-size: 9pt; color: #999; text-align: center; }}
  @media print {{ body {{ background: #fff; }} .container {{ box-shadow: none; padding: 20px; }} .page-content {{ break-inside: avoid; }} }}
</style>
</head>
<body>
<div class="container">
  <!-- Header -->
  <h1>{title}</h1>
  <div class="meta-row">
    <span><span class="meta-label">Authors:</span> {_escape_html(authors_str)}</span>
    <span><span class="meta-label">Year:</span> {year_str}</span>
    <span><span class="meta-label">Language:</span> <span class="badge badge-lang">{lang}</span></span>
    <span><span class="meta-label">Pages:</span> {pages}</span>
    {f'<span class="badge badge-starred">{starred_badge}</span>' if data["starred"] else ''}
    <span><span class="meta-label">Tags:</span> {_escape_html(tags_str)}</span>
    <span><span class="meta-label">Status:</span> <span class="badge badge-status">{data["status"]}</span></span>
  </div>

  <!-- AI Summary -->
  {summary_html}

  <!-- Personal Notes -->
  {notes_html}

  <!-- Content by Page -->
  <section class="section">
    <h2>📖 Full Text Content</h2>
    {content_html}
  </section>

  <!-- Citation -->
  <section class="section citation-section">
    <h2>🔖 Citation (APA 7th)</h2>
    <div class="citation-text">{citation}</div>
  </section>

  <!-- Footer -->
  <div class="footer">
    Exported from ResearchMind VN on {today}
  </div>
</div>
</body>
</html>"""

    return HTMLResponse(content=html, status_code=200)


# ─── DOCX Export ────────────────────────────────────────────────

@router.get("/{paper_id}/export/docx")
async def export_paper_docx(paper_id: str, db: Session = Depends(_get_db), request: Request = None):
    """Export full paper content as a Word document (.docx)."""
    lang = get_language(request) if request else "vi"
    data = _get_paper_data(paper_id, db)
    if not data:
        raise HTTPException(status_code=404, detail="Paper not found")

    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail=t("export.docx_not_installed", lang),
        )

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # ── Title ──
    title_para = doc.add_heading(data["title"], level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata Table ──
    authors_str = ", ".join(data["authors"]) if data["authors"] else "Unknown"
    tags_str = ", ".join(data["tags"]) if data["tags"] else "—"

    meta_table = doc.add_table(rows=7, cols=2, style="Light Shading Accent 1")
    meta_data = [
        ("Authors", authors_str),
        ("Year", str(data["year"]) if data["year"] else "N/A"),
        ("DOI", data["doi"] or "—"),
        ("Language", data["language"].upper()),
        ("Pages", str(data["page_count"])),
        ("Tags", tags_str),
        ("Status", data["status"]),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        # Bold the label
        for paragraph in meta_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()  # spacer

    # ── AI Summary ──
    if data["auto_summary"]:
        doc.add_heading("🧠 AI Summary", level=2)
        # Strip markdown formatting for clean docx
        summary_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", data["auto_summary"])
        summary_clean = re.sub(r"\*([^*]+)\*", r"\1", summary_clean)
        summary_clean = re.sub(r"^###\s*", "", summary_clean, flags=re.MULTILINE)
        summary_clean = re.sub(r"^\s*[\-\*]\s+", "• ", summary_clean, flags=re.MULTILINE)
        doc.add_paragraph(summary_clean)

    # ── Personal Notes ──
    if data["notes"]:
        doc.add_heading("📝 Personal Notes", level=2)
        doc.add_paragraph(data["notes"])

    # ── Content by Page ──
    doc.add_heading("📖 Full Text Content", level=2)
    for page_num in sorted(data["chunks_by_page"].keys()):
        chunks_text = data["chunks_by_page"][page_num]
        doc.add_heading(f"Page {page_num}", level=3)
        for chunk_text in chunks_text:
            # Split into paragraphs for readability
            for para_text in chunk_text.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Cm(1.27)
        doc.add_paragraph()  # spacing between pages

    # ── Citation ──
    doc.add_heading("🔖 Citation (APA 7th)", level=2)
    if data["authors"]:
        if len(data["authors"]) == 1:
            cite_author = data["authors"][0]
        elif len(data["authors"]) == 2:
            cite_author = f"{data['authors'][0]} & {data['authors'][1]}"
        else:
            cite_author = f"{data['authors'][0]} et al."
    else:
        cite_author = "Unknown"
    year_str = str(data["year"]) if data["year"] else "n.d."
    citation_text = f"{cite_author} ({year_str}). {data['title']}."
    if data["doi"]:
        citation_text += f" https://doi.org/{data['doi']}"
    doc.add_paragraph(citation_text)

    # ── Footer ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Exported from ResearchMind VN on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Set margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Save to BytesIO and return ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_filename = data["filename"].replace(".pdf", "")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_filename}.docx"',
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    )


# ─── Synthesis Export ──────────────────────────────────────────

from pydantic import BaseModel

class SynthesisExportRequest(BaseModel):
    title: str
    content: str
    format: str  # "docx" or "html" or "markdown"

@router.post("/export/synthesis")
async def export_synthesis(
    req: SynthesisExportRequest,
    request: Request,
):
    """Export arbitrary synthesis markdown content as a DOCX or HTML or Markdown file."""
    lang = get_language(request)
    title = req.title
    content = req.content
    fmt = req.format.lower()

    if fmt == "markdown" or fmt == "md":
        buf = io.BytesIO(content.encode("utf-8"))
        buf.seek(0)
        safe_title = re.sub(r"[^\w\-]", "_", title)
        return StreamingResponse(
            buf,
            media_type="text/markdown",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.md"',
                "Content-Type": "text/markdown",
            },
        )

    elif fmt == "html":
        # Convert Markdown to HTML
        html_body = _md_to_html(content)
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{_escape_html(title)}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            color: #1e293b;
            background: #f8fafc;
        }}
        /* ── Theme Toolbar ── */
        .toolbar {{
            display: flex;
            align-items: center;
            justify-content: flex-end;
            gap: 8px;
            margin-bottom: 16px;
            padding: 8px 12px;
            background: #f1f5f9;
            border-radius: 10px;
            border: 1px solid #e2e8f0;
        }}
        .toolbar-label {{
            font-size: 11px;
            font-weight: 600;
            color: #64748b;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-right: auto;
        }}
        .theme-btn {{
            padding: 4px 10px;
            border: 1px solid #e2e8f0;
            border-radius: 6px;
            background: #ffffff;
            color: #475569;
            font-size: 11px;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.15s ease;
        }}
        .theme-btn:hover {{
            background: #f1f5f9;
            border-color: #cbd5e1;
        }}
        .theme-btn.active {{
            background: #6366f1;
            border-color: #6366f1;
            color: #ffffff;
        }}
        .theme-btn-dark.active {{
            background: #8b5cf6;
            border-color: #8b5cf6;
        }}
        .theme-btn::before {{
            content: "";
            display: inline-block;
            width: 8px;
            height: 8px;
            border-radius: 50%;
            margin-right: 5px;
            vertical-align: middle;
        }}
        .theme-btn[data-theme="github"]::before {{ background: #f5f5f5; border: 1px solid #ddd; }}
        .theme-btn[data-theme="github-dark"]::before {{ background: #0d1117; }}
        .theme-btn[data-theme="monokai"]::before {{ background: #272822; }}
        .theme-btn[data-theme="dracula"]::before {{ background: #282a36; }}
        .theme-btn[data-theme="atom-one-dark"]::before {{ background: #1e1e1e; }}
        .theme-btn[data-theme="nord"]::before {{ background: #2e3440; }}
        .theme-btn[data-theme="vs2015"]::before {{ background: #1e1e1e; }}
        .theme-btn[data-theme="tomorrow"]::before {{ background: #ffffff; border: 1px solid #ddd; }}
        .container {{
            border-bottom: 2px solid #8b5cf6;
            padding-bottom: 12px;
            color: #0f172a;
            margin-top: 0;
        }}
        h2 {{
            color: #1e293b;
            margin-top: 32px;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 8px;
        }}
        h3 {{
            color: #334155;
        }}
        p {{
            margin-bottom: 1.25em;
        }}
        ul, ol {{
            padding-left: 20px;
            margin-bottom: 1.25em;
        }}
        li {{
            margin-bottom: 0.5em;
        }}
        code {{
            background: #f1f5f9;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 0.9em;
            font-family: SFMono-Regular, Consolas, "Liberation Mono", Menlo, monospace;
        }}
        pre {{
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 16px 20px;
            overflow-x: auto;
            margin: 0 0 1.25em 0;
        }}
        pre code {{
            padding: 0;
            border-radius: 0;
            font-size: 0.82em;
            line-height: 1.5;
            display: block;
            white-space: pre;
            background: none !important;
        }}
        /* Hide language label from highlight.js */
        pre code.hljs::before {{
            display: none;
        }}
        blockquote {{
            border-left: 4px solid #8b5cf6;
            background: #f8fafc;
            margin: 0 0 1.25em 0;
            padding: 12px 20px;
            color: #475569;
            font-style: italic;
        }}
        hr {{
            border: none;
            border-top: 1px solid #e2e8f0;
            margin: 32px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 0 0 1.25em 0;
            font-size: 0.9em;
        }}
        table thead {{
            background: #6366f1;
            color: #ffffff;
        }}
        table th {{
            padding: 10px 12px;
            text-align: center;
            font-weight: 600;
        }}
        table td {{
            padding: 8px 12px;
            text-align: center;
            border-bottom: 1px solid #e2e8f0;
            color: #334155;
        }}
        table tbody tr:nth-child(even) {{
            background: #f8fafc;
        }}
        table tbody tr:hover {{
            background: #f1f5f9;
        }}
        .footer {{
            margin-top: 60px;
            border-top: 1px solid #e2e8f0;
            padding-top: 20px;
            font-size: 0.85em;
            color: #64748b;
            text-align: center;
        }}
    </style>
    <!-- highlight.js for syntax highlighting -->
    <link id="hljs-theme" rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github.min.css">
</head>
<body>
    <!-- Theme Toolbar -->
    <div class="toolbar">
        <span class="toolbar-label">🎨 Theme</span>
        <button class="theme-btn active" data-theme="github" onclick="setTheme('github')" title="GitHub Light">Light</button>
        <button class="theme-btn theme-btn-dark" data-theme="github-dark" onclick="setTheme('github-dark')" title="GitHub Dark">GitHub</button>
        <button class="theme-btn theme-btn-dark" data-theme="monokai" onclick="setTheme('monokai')" title="Monokai">Monokai</button>
        <button class="theme-btn theme-btn-dark" data-theme="dracula" onclick="setTheme('dracula')" title="Dracula">Dracula</button>
        <button class="theme-btn theme-btn-dark" data-theme="atom-one-dark" onclick="setTheme('atom-one-dark')" title="Atom One Dark">Atom</button>
        <button class="theme-btn theme-btn-dark" data-theme="nord" onclick="setTheme('nord')" title="Nord">Nord</button>
        <button class="theme-btn theme-btn-dark" data-theme="vs2015" onclick="setTheme('vs2015')" title="VS2015">VS</button>
        <button class="theme-btn" data-theme="tomorrow" onclick="setTheme('tomorrow')" title="Tomorrow">Tomorrow</button>
    </div>
    <div class="container">
        <h1>{_escape_html(title)}</h1>
        {html_body}
        <div class="footer">
            Exported from ResearchMind VN on {datetime.now().strftime('%Y-%m-%d %H:%M')}
        </div>
    </div>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/highlight.min.js"></script>
    <script>
        const THEMES = {{
            'github':     'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github.min.css',
            'github-dark':'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github-dark.min.css',
            'monokai':    'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/monokai.min.css',
            'dracula':    'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/dracula.min.css',
            'atom-one-dark':'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/atom-one-dark.min.css',
            'nord':       'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/nord.min.css',
            'vs2015':     'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/vs2015.min.css',
            'tomorrow':   'https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/tomorrow.min.css',
        }};

        function applyTheme(name) {{
            const link = document.getElementById('hljs-theme');
            if (!link || !THEMES[name]) return;
            link.onload = function() {{
                document.querySelectorAll('pre code').forEach(function(block) {{
                    hljs.highlightElement(block);
                }});
            }};
            link.href = THEMES[name];
            document.querySelectorAll('.theme-btn').forEach(function(btn) {{
                btn.classList.toggle('active', btn.getAttribute('data-theme') === name);
            }});
        }}

        function setTheme(name) {{
            applyTheme(name);
            localStorage.setItem('hljs-theme', name);
            // User made an explicit choice — detach system preference listener
            try {{ window.matchMedia('(prefers-color-scheme: dark)').removeEventListener('change', autoSwitchTheme); }} catch(e) {{}}
        }}

        function autoSwitchTheme(e) {{
            if (!localStorage.getItem('hljs-theme')) {{
                applyTheme(e.matches ? 'github-dark' : 'github');
            }}
        }}

        function getSystemTheme() {{
            return window.matchMedia('(prefers-color-scheme: dark)').matches ? 'github-dark' : 'github';
        }}

        // Initialize: saved preference > system dark mode > default (github)
        (function() {{
            const saved = localStorage.getItem('hljs-theme');
            if (saved && THEMES[saved]) {{
                applyTheme(saved);  // user's explicit choice
            }} else {{
                applyTheme(getSystemTheme());  // auto-detect, don't save
                window.matchMedia('(prefers-color-scheme: dark)').addEventListener('change', autoSwitchTheme);
            }}
        }})();
    </script>
</body>
</html>"""
        buf = io.BytesIO(html_content.encode("utf-8"))
        buf.seek(0)
        safe_title = re.sub(r"[^\w\-]", "_", title)
        return StreamingResponse(
            buf,
            media_type="text/html",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.html"',
                "Content-Type": "text/html",
            },
        )

    elif fmt == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail=t("export.fpdf_not_installed", lang),
            )

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)

        # ════════════════════════════════════════════════════════════
        # COVER PAGE
        # ════════════════════════════════════════════════════════════
        pdf.add_page()
        pdf.alias_nb_pages()

        # Spacer to push title to vertical center
        pdf.ln(50)

        # Decorative top line (purple)
        pdf.set_draw_color(99, 102, 241)
        pdf.set_line_width(1.2)
        mid_x = pdf.w / 2
        pdf.line(mid_x - 30, pdf.get_y(), mid_x + 30, pdf.get_y())
        pdf.ln(14)

        # Title (large, centered, dark)
        pdf.set_font("Helvetica", "B", 28)
        pdf.set_text_color(15, 23, 42)
        title_text = title.replace("_", " ").replace("Report", "Report")
        pdf.multi_cell(0, 14, title_text, align="C")
        pdf.ln(6)

        # Subtitle
        pdf.set_font("Helvetica", "", 12)
        pdf.set_text_color(100, 116, 139)
        if "Literature" in title_text:
            pdf.cell(0, 8, "Literature Review Report", align="C")
        elif "Critique" in title_text:
            pdf.cell(0, 8, "Paper Critique Report", align="C")
        elif "Debate" in title_text:
            pdf.cell(0, 8, "AI Debate Transcript", align="C")
        else:
            pdf.cell(0, 8, "Research Synthesis Report", align="C")
        pdf.ln(14)

        # Decorative bottom line (purple)
        pdf.set_draw_color(99, 102, 241)
        pdf.set_line_width(1.2)
        y_now = pdf.get_y()
        pdf.line(mid_x - 30, y_now, mid_x + 30, y_now)
        pdf.ln(40)

        # Metadata at bottom
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(148, 163, 184)
        pdf.cell(0, 6, f"Generated by ResearchMind VN", align="C")
        pdf.ln(5)
        pdf.cell(0, 6, f"{datetime.now().strftime('%B %d, %Y at %H:%M')}", align="C")
        pdf.ln(5)
        pdf.cell(0, 6, "ResearchMind VN — AI-Powered Research Assistant", align="C")

        # ── Page break after cover ──
        pdf.add_page()

        # ════════════════════════════════════════════════════════════
        # CONTENT PAGES — consumes shared _parse_md_events()
        # ════════════════════════════════════════════════════════════
        pdf.set_text_color(30, 30, 30)

        def _pdf_strip(text: str) -> str:
            return re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"\*(.+?)\*", r"\1", text))

        def _pdf_code_block(code_lines: list[str], lang: str):
            line_h = 5
            x0 = pdf.l_margin
            y0 = pdf.get_y()
            code_w = pdf.w - pdf.l_margin - pdf.r_margin
            padding = 3
            lang_h = 4 if lang else 0
            x_start = x0 + padding
            y_start = y0 + padding
            total_lines = max(len(code_lines), 1)
            block_h = total_lines * line_h + 2 * padding + 2 + lang_h
            if y0 + block_h > pdf.h - 20:
                pdf.add_page()
                y0 = pdf.get_y()
                y_start = y0 + padding
            pdf.set_fill_color(241, 245, 249)
            pdf.set_draw_color(203, 213, 225)
            pdf.rect(x0, y0, code_w, block_h, style="DF")
            if lang:
                pdf.set_font("Helvetica", "I", 7)
                pdf.set_text_color(148, 163, 184)
                pdf.set_xy(x_start + 1, y_start)
                pdf.cell(0, 4, lang)
                y_start += lang_h
            pdf.set_font("Courier", "", 8)
            pdf.set_text_color(30, 41, 59)
            pdf.set_xy(x_start, y_start)
            for cl in code_lines:
                pdf.cell(0, line_h, cl.replace("\t", "    "))
                pdf.set_x(x_start)
            pdf.set_y(y_start + len(code_lines) * line_h + padding)
            pdf.ln(2)

        def _pdf_table(rows: list[list[str]]):
            if len(rows) < 2:
                return
            header_row = rows[0]
            data_rows = rows[2:] if len(rows) > 2 else []
            n_cols = len(header_row)
            if n_cols == 0:
                return
            col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / n_cols
            cell_h = 7
            y0 = pdf.get_y()
            if y0 + (cell_h * (1 + len(data_rows)) + 2) > pdf.h - 20:
                pdf.add_page()
                y0 = pdf.get_y()
            x0 = pdf.l_margin
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(99, 102, 241)
            pdf.set_text_color(255, 255, 255)
            pdf.set_draw_color(99, 102, 241)
            pdf.set_xy(x0, y0)
            for hcell in header_row:
                pdf.cell(col_w, cell_h, hcell.strip(), border=1, fill=True, align="C")
            pdf.ln(cell_h)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 30, 30)
            pdf.set_draw_color(203, 213, 225)
            for ri, row in enumerate(data_rows):
                pdf.set_fill_color(248, 250, 252) if ri % 2 == 0 else pdf.set_fill_color(255, 255, 255)
                pdf.set_xy(x0, pdf.get_y())
                for dcell in row:
                    pdf.cell(col_w, cell_h, dcell.strip(), border=1, fill=True, align="C")
                pdf.ln(cell_h)
            pdf.ln(3)

        for event, data in _parse_md_events(content):
            if event == "code_block":
                _pdf_code_block(data[0], data[1])
            elif event == "table":
                _pdf_table(data[0])
            elif event == "heading1":
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(15, 23, 42)
                pdf.multi_cell(0, 10, _pdf_strip(data[0]))
                pdf.ln(2)
            elif event == "heading2":
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(30, 41, 59)
                pdf.multi_cell(0, 8, _pdf_strip(data[0]))
                pdf.ln(1)
            elif event == "heading3":
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(51, 65, 85)
                pdf.multi_cell(0, 7, _pdf_strip(data[0]))
                pdf.ln(1)
            elif event == "hr":
                pdf.set_draw_color(200, 200, 200)
                pdf.set_line_width(0.3)
                pdf.line(pdf.l_margin, pdf.get_y() + 2, pdf.w - pdf.r_margin, pdf.get_y() + 2)
                pdf.ln(6)
            elif event == "bullet_list":
                for item in data[0]:
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(30, 30, 30)
                    pdf.cell(6, 6, chr(8226))
                    pdf.multi_cell(0, 6, _pdf_strip(item))
            elif event == "numbered_list":
                for idx, item in enumerate(data[0], 1):
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(30, 30, 30)
                    pdf.multi_cell(0, 6, f"{idx}. {_pdf_strip(item)}")
            elif event == "blockquote":
                for line in data[0]:
                    pdf.set_x(pdf.l_margin + 4)
                    pdf.set_font("Helvetica", "I", 10)
                    pdf.set_text_color(71, 85, 105)
                    y_start = pdf.get_y()
                    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 8, 6, line if line else " ")
                    y_end = pdf.get_y()
                    pdf.set_draw_color(139, 92, 246)
                    pdf.set_line_width(0.8)
                    pdf.line(pdf.l_margin, y_start, pdf.l_margin, y_end)
                    pdf.set_text_color(30, 30, 30)
            elif event == "paragraph":
                for line in data[0]:
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(30, 30, 30)
                    pdf.multi_cell(0, 6, _pdf_strip(line))

        # ── Footer with page numbers ──
        pdf.footer = lambda: _pdf_footer(pdf)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)

        safe_title = re.sub(r"[^\w\-]", "_", title)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.pdf"',
                "Content-Type": "application/pdf",
            },
        )

    elif fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail=t("export.docx_not_installed", lang),
            )

        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        # Title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # spacer

        # Parse and write markdown to docx (consumes shared _parse_md_events)
        def _docx_code_block(code_lines: list[str], lang: str):
            if lang:
                lp = doc.add_paragraph()
                lr = lp.add_run(lang)
                lr.font.size = Pt(8)
                lr.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
                lr.italic = True
                lp.paragraph_format.space_after = Pt(2)
                lp.paragraph_format.space_before = Pt(4)
            for cl in code_lines:
                cp = doc.add_paragraph()
                cr = cp.add_run(cl.replace("\t", "    ") if cl else " ")
                cr.font.name = "Consolas"
                cr.font.size = Pt(8.5)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.line_spacing = 1.15
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'f1f5f9')
                shd.set(qn('w:val'), 'clear')
                cp.paragraph_format.element.get_or_add_pPr().append(shd)

        def _docx_table(rows: list[list[str]]):
            if len(rows) < 2:
                return
            header_row = rows[0]
            data_rows = rows[2:] if len(rows) > 2 else []
            if not header_row:
                return
            table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
            table.style = 'Table Grid'
            for ci, hcell in enumerate(header_row):
                cell = table.rows[0].cells[ci]
                cell.text = hcell.strip()
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '6366f1')
                shd.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shd)
            alt = ['f8fafc', 'ffffff']
            for ri, row in enumerate(data_rows):
                for ci, dcell in enumerate(row):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = dcell.strip()
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), alt[ri % 2])
                    shd.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shd)
            doc.add_paragraph()

        for event, data in _parse_md_events(content):
            if event == "code_block":
                _docx_code_block(data[0], data[1])
            elif event == "table":
                _docx_table(data[0])
            elif event == "heading1":
                doc.add_heading(data[0], level=1)
            elif event == "heading2":
                doc.add_heading(data[0], level=2)
            elif event == "heading3":
                doc.add_heading(data[0], level=3)
            elif event == "hr":
                p = doc.add_paragraph()
                pPr = p.paragraph_format.element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '4')
                bottom.set(qn('w:color'), 'cbd5e1')
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif event == "bullet_list":
                for item in data[0]:
                    p = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(p, item)
            elif event == "numbered_list":
                for item in data[0]:
                    p = doc.add_paragraph(style='List Number')
                    _add_formatted_text(p, item)
            elif event == "blockquote":
                for line in data[0]:
                    p = doc.add_paragraph(style='Quote')
                    _add_formatted_text(p, line)
            elif event == "paragraph":
                for line in data[0]:
                    p = doc.add_paragraph()
                    _add_formatted_text(p, line)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Exported from ResearchMind VN on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_title = re.sub(r"[^\w\-]", "_", title)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.docx"',
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'docx', 'html', or 'markdown'")


def _pdf_footer(pdf):
    """Add page number to PDF footer."""
    pdf.set_y(-15)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 10, f"Page {pdf.page_no()}/{{nb}}", align="C")


def _add_formatted_text(paragraph, text: str):
    # Split text by bold markers (**)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            bold_text = part[2:-2]
            # Further split by italic markers (*)
            subparts = re.split(r"(\*.*?\*)", bold_text)
            for subpart in subparts:
                if subpart.startswith("*") and subpart.endswith("*"):
                    run = paragraph.add_run(subpart[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(subpart)
                    run.bold = True
        else:
            subparts = re.split(r"(\*.*?\*)", part)
            for subpart in subparts:
                if subpart.startswith("*") and subpart.endswith("*"):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(subpart)
