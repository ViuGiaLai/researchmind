"""
Test script: Generate a DOCX from sample markdown content with code blocks and tables.
Uses the same logic as the synthesis export endpoint's DOCX handler.
"""

import io
import re
import sys
import subprocess

# ─── First, ensure python-docx is installed ─────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    print("✅ python-docx is already installed")
except ImportError:
    print("📦 Installing python-docx...")
    result = subprocess.run(
        [sys.executable, "-m", "pip", "install", "python-docx"],
        capture_output=True, text=True
    )
    if result.returncode == 0:
        print("✅ python-docx installed successfully")
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    else:
        print(f"❌ Failed to install python-docx: {result.stderr}")
        sys.exit(1)


# ─── Formatted text helper (same as export.py) ─────────────────
def _add_formatted_text(paragraph, text: str):
    """Add formatted text with bold/italic support."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            bold_text = part[2:-2]
            subparts = re.split(r"(\*.*?\*)", bold_text)
            for subpart in subparts:
                if subpart.startswith("*") and subpart.endswith("*"):
                    run = paragraph.add_run(subpart[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(subpart)
                    run.bold = True
        else:
            subparts = re.split(r"(\*.*?\*)", part)
            for subpart in subparts:
                if subpart.startswith("*") and subpart.endswith("*"):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(subpart)


# ─── Generate DOCX from Markdown (same logic as export.py) ─────
def generate_docx(title: str, content: str, output_path: str):
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Title
    title_para = doc.add_heading(title.replace("_", " "), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Multi-line state machine parser ──
    lines = content.split("\n")
    i = 0
    n = len(lines)

    def _add_code_block(doc_obj, code_lines: list[str], lang: str):
        if lang:
            lp = doc_obj.add_paragraph()
            lr = lp.add_run(lang)
            lr.font.size = Pt(8)
            lr.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
            lr.italic = True
            lp.paragraph_format.space_after = Pt(2)
            lp.paragraph_format.space_before = Pt(4)

        for code_line in code_lines:
            cp = doc_obj.add_paragraph()
            cr = cp.add_run(code_line.replace("\t", "    ") if code_line else " ")
            cr.font.name = "Consolas"
            cr.font.size = Pt(8.5)
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.line_spacing = 1.15
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'f1f5f9')
            shading_elm.set(qn('w:val'), 'clear')
            cp.paragraph_format.element.get_or_add_pPr().append(shading_elm)

    def _add_docx_table(doc_obj, rows: list[list[str]]):
        if len(rows) < 2:
            return
        header_row = rows[0]
        data_rows = rows[2:] if len(rows) > 2 else []
        if not header_row:
            return

        table = doc_obj.add_table(rows=1 + len(data_rows), cols=len(header_row))
        table.style = 'Table Grid'

        # Header
        for ci, hcell in enumerate(header_row):
            cell = table.rows[0].cells[ci]
            cell.text = hcell.strip()
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '6366f1')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        alt_colors = ['f8fafc', 'ffffff']
        for ri, row in enumerate(data_rows):
            row_idx = ri + 1
            for ci, dcell in enumerate(row):
                cell = table.rows[row_idx].cells[ci]
                cell.text = dcell.strip()
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                shade = OxmlElement('w:shd')
                shade.set(qn('w:fill'), alt_colors[ri % 2])
                shade.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shade)

        doc_obj.add_paragraph()

    while i < n:
        stripped = lines[i].strip()

        # Code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            _add_code_block(doc, code_lines, lang)
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().split("|")]
                cells = [c for c in cells if c]
                table_rows.append(cells)
                i += 1
            _add_docx_table(doc, table_rows)
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith("---"):
            p = doc.add_paragraph()
            pPr = p.paragraph_format.element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), 'cbd5e1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
            content_text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, content_text)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(style='Quote')
            _add_formatted_text(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1

    # Footer
    from datetime import datetime
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Exported from ResearchMind VN on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    doc.save(output_path)
    return output_path


# ─── Test with rich sample markdown ────────────────────────────
if __name__ == "__main__":
    sample_markdown = """# Synthesis Report
*Model: gemini-2.0-flash*

## Overview
This is a **sample synthesis report** testing the enhanced *DOCX export* with code blocks, tables, and horizontal rules.

### Key Findings
- Research shows significant improvement in **model performance**
- The proposed method achieves **94.2% accuracy** on benchmark datasets
- Multi-dataset evaluation with comprehensive metrics

> This finding strongly supports the main hypothesis and aligns with prior work in the field.

---

## Code Implementation

The main training loop is implemented as follows:

```python
import torch
import torch.nn as nn

class ResearchModel(nn.Module):
    def __init__(self, vocab_size, hidden_dim=768):
        super().__init__()
        self.embed = nn.Embedding(vocab_size, hidden_dim)
        self.transformer = nn.TransformerEncoder(
            nn.TransformerEncoderLayer(hidden_dim, nhead=12),
            num_layers=6
        )

    def forward(self, x):
        return self.transformer(self.embed(x))
```

To run training:

```bash
python train.py --epochs 10 --batch-size 32 --lr 2e-5
```

---

## Experimental Results

### Benchmark Comparison
| Model | Accuracy | F1 Score | Precision | Recall |
|-------|----------|----------|-----------|--------|
| Ours | 94.2% | 0.937 | 0.945 | 0.931 |
| Baseline A | 89.1% | 0.882 | 0.891 | 0.874 |
| Baseline B | 91.5% | 0.908 | 0.914 | 0.902 |
| SOTA 2024 | 93.8% | 0.931 | 0.938 | 0.925 |

Our method outperforms all baselines across every metric with a **2.7% improvement** over the previous state-of-the-art.

### Ablation Study
| Component | Accuracy | Delta |
|-----------|----------|-------|
| Full model | 94.2% | — |
| w/o attention | 88.3% | -5.9% |
| w/o pretraining | 90.1% | -4.1% |
| w/o data augmentation | 92.7% | -1.5% |

> "The attention mechanism contributes most significantly to the overall performance."

---

## Conclusion
This work presents a novel approach with the following key contributions:

1. **Novel architecture** combining transformers with retrieval augmentation
2. **Comprehensive evaluation** across 5 benchmark datasets
3. **State-of-the-art results** with 94.2% accuracy
4. **Open-source release** of code and pre-trained models

Future work includes:
* Scaling to larger model sizes
* Multi-modal extension with vision encoders
* Real-time deployment optimization

---

*This test document validates code blocks, tables, and horizontal rules in DOCX export.*
"""

    output_path = "test_synthesis_report.docx"
    try:
        result = generate_docx("Enhanced_Synthesis_Report_Test", sample_markdown, output_path)
        size = __import__('os').path.getsize(output_path)
        print(f"\n✅ DOCX generated successfully: {output_path}")
        print(f"   File size: {size} bytes ({size/1024:.1f} KB)")
        print(f"   Location: {__import__('os').path.abspath(output_path)}")
    except Exception as e:
        print(f"\n❌ DOCX generation failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
